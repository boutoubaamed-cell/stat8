"""
تطبيق Streamlit للتحليل الإحصائي المتقدم للاستبيانات - نسخة محسنة مع دعم RTL
"""

import streamlit as st
import pandas as pd
import numpy as np
import io
import warnings
import datetime
from io import BytesIO
from datetime import datetime
import os
import traceback
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart

warnings.filterwarnings('ignore')

# ============================================
# Initialize library availability flags
# ============================================
STATS_AVAILABLE = False
PLOTS_AVAILABLE = False
DOCX_AVAILABLE = False
EXCEL_AVAILABLE = False

# ============================================
# Page configuration - MUST BE FIRST STREAMLIT COMMAND
# ============================================
st.set_page_config(
    page_title="التحليل الإحصائي المتقدم",
    page_icon="📊",
    layout="wide",
    initial_sidebar_state="expanded"
)

# ============================================
# Try to import optional libraries with graceful fallback
# ============================================
# Try to import matplotlib and seaborn
try:
    import matplotlib
    matplotlib.use('Agg')  # Use non-interactive backend
    import matplotlib.pyplot as plt
    import seaborn as sns
    # إعدادات الخط العربي للمخططات
    plt.rcParams['font.family'] = 'Arial'
    plt.rcParams['axes.unicode_minus'] = False
    PLOTS_AVAILABLE = True
except ImportError:
    # Define placeholder functions if matplotlib is not available
    class plt:
        @staticmethod
        def figure(*args, **kwargs):
            return None
        @staticmethod
        def close(*args, **kwargs):
            pass
    class sns:
        @staticmethod
        def heatmap(*args, **kwargs):
            pass

# Try to import openpyxl for Excel support
try:
    import openpyxl
    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False

# Try to import scipy and statsmodels
try:
    from scipy import stats
    from scipy.stats import (
        pearsonr, spearmanr, kendalltau,
        ttest_ind, ttest_rel, f_oneway,
        chi2_contingency, mannwhitneyu, wilcoxon,
        kruskal, friedmanchisquare,
        shapiro, normaltest, anderson, levene, bartlett
    )
    import statsmodels.api as sm
    from statsmodels.stats.multicomp import pairwise_tukeyhsd
    STATS_AVAILABLE = True
except ImportError:
    pass

# Try to import docx
try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    pass

# ============================================
# CSS for RTL support and better styling
# ============================================
st.markdown("""
<style>
    /* RTL Support */
    .stApp {
        direction: rtl;
        text-align: right;
    }
    .stMarkdown, .stText, .stLatex, .stTitle, .stHeader, .stSubheader {
        direction: rtl;
        text-align: right;
    }
    .stDataFrame {
        direction: rtl;
        text-align: right;
    }
    .stSelectbox, .stMultiselect {
        direction: rtl;
        text-align: right;
    }
    .stButton > button {
        direction: rtl;
        width: 100%;
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        color: white;
        border: none;
        padding: 10px 20px;
        border-radius: 5px;
        font-weight: bold;
    }
    .stButton > button:hover {
        background: linear-gradient(135deg, #764ba2 0%, #667eea 100%);
    }
    
    /* File uploader styling */
    .stFileUploader {
        direction: rtl;
        text-align: center;
        border: 2px dashed #1e3c72;
        border-radius: 10px;
        padding: 20px;
        background-color: #f8f9fa;
    }
    .uploadedFile {
        direction: rtl;
        text-align: center;
        padding: 10px;
        margin: 10px 0;
        background: linear-gradient(135deg, #43a047 0%, #2e7d32 100%);
        color: white;
        border-radius: 5px;
    }
    
    /* Result cards styling */
    .result-card {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        padding: 20px;
        border-radius: 10px;
        color: white;
        margin: 10px 0;
        box-shadow: 0 4px 6px rgba(0,0,0,0.1);
    }
    
    /* Likert trend indicators */
    .trend-very-high {
        background: linear-gradient(135deg, #1b5e20 0%, #2e7d32 100%);
        color: white;
        padding: 15px;
        border-radius: 10px;
        text-align: center;
        font-weight: bold;
        margin: 5px 0;
    }
    .trend-high {
        background: linear-gradient(135deg, #2e7d32 0%, #388e3c 100%);
        color: white;
        padding: 15px;
        border-radius: 10px;
        text-align: center;
        font-weight: bold;
        margin: 5px 0;
    }
    .trend-moderate {
        background: linear-gradient(135deg, #f9a825 0%, #f57f17 100%);
        color: white;
        padding: 15px;
        border-radius: 10px;
        text-align: center;
        font-weight: bold;
        margin: 5px 0;
    }
    .trend-low {
        background: linear-gradient(135deg, #c62828 0%, #b71c1c 100%);
        color: white;
        padding: 15px;
        border-radius: 10px;
        text-align: center;
        font-weight: bold;
        margin: 5px 0;
    }
    .trend-very-low {
        background: linear-gradient(135deg, #b71c1c 0%, #8e0000 100%);
        color: white;
        padding: 15px;
        border-radius: 10px;
        text-align: center;
        font-weight: bold;
        margin: 5px 0;
    }
    
    /* Contact form styling */
    .contact-info {
        background: linear-gradient(135deg, #1e3c72 0%, #2a5298 100%);
        padding: 30px;
        border-radius: 15px;
        color: white;
        margin: 20px 0;
        text-align: center;
    }
    .contact-info h3 {
        color: white;
        margin-bottom: 20px;
    }
    .contact-info p {
        color: #e0e0e0;
        font-size: 1.3em;
        line-height: 1.8;
    }
    
    /* Section headers */
    .section-header {
        background: linear-gradient(135deg, #1e3c72 0%, #2a5298 100%);
        color: white;
        padding: 10px 15px;
        border-radius: 5px;
        margin: 20px 0 10px 0;
    }
    
    /* Plot container - LTR for charts */
    .plot-container {
        direction: ltr !important;
        text-align: center !important;
    }
    .js-plotly-plot, .plot-container, .stPlotlyChart {
        direction: ltr !important;
    }
    
    /* Installation instructions */
    .install-box {
        background: linear-gradient(135deg, #ff6b6b 0%, #c92a2a 100%);
        color: white;
        padding: 15px;
        border-radius: 10px;
        text-align: center;
        margin: 10px 0;
    }
    .install-code {
        background: #2d2d2d;
        color: #ffd700;
        padding: 10px;
        border-radius: 5px;
        direction: ltr;
        text-align: left;
        font-family: monospace;
        margin: 10px 0;
    }
</style>
""", unsafe_allow_html=True)

# ============================================
# Header
# ============================================
st.markdown("""
<div style='text-align: center; padding: 30px; background: linear-gradient(135deg, #1e3c72 0%, #2a5298 100%); border-radius: 15px; margin-bottom: 25px;'>
    <h1 style='color: white; font-size: 2.5em; margin-bottom: 10px;'>📊 منصة التحليل الإحصائي المتقدم </h1>
    <p style='color: #e0e0e0; font-size: 1.2em;'>اختبارات معلمية ولا معلمية - تحليل العوامل - رسوم بيانية تفاعلية</p>
    <p style='color: #b8c7e0; font-size: 1em; margin-top: 10px;'>بوابتك المتكاملة لإنجاز البحوث والدراسات الكمية باحترافية </p>
</div>
""", unsafe_allow_html=True)

# ============================================
# Session state initialization
# ============================================
if 'data_loaded' not in st.session_state:
    st.session_state.data_loaded = False
if 'df' not in st.session_state:
    st.session_state.df = None
if 'social_vars' not in st.session_state:
    st.session_state.social_vars = []
if 'factors' not in st.session_state:
    st.session_state.factors = {}
if 'show_results' not in st.session_state:
    st.session_state.show_results = False
if 'factor_trends' not in st.session_state:
    st.session_state.factor_trends = {}
if 'uploaded_filename' not in st.session_state:
    st.session_state.uploaded_filename = None
if 'contact_messages' not in st.session_state:
    st.session_state.contact_messages = []

# ============================================
# Email configuration
# ============================================
CONTACT_EMAIL = "boutoubaamed@gmail.com"

# ============================================
# File upload functions with improved error handling
# ============================================
@st.cache_data
def load_csv_with_encoding(file):
    """Try to load CSV with different encodings"""
    encodings = ['utf-8', 'cp1256', 'ISO-8859-1', 'latin1', 'utf-16']

    for encoding in encodings:
        try:
            file.seek(0)
            df = pd.read_csv(file, encoding=encoding)
            return df, encoding
        except:
            continue

    try:
        file.seek(0)
        df = pd.read_csv(file)
        return df, 'default'
    except Exception as e:
        raise Exception(f"تعذر قراءة ملف CSV: {str(e)}")

def load_excel_file(file):
    """Load Excel file with error handling and dependency check"""
    if not EXCEL_AVAILABLE:
        raise ImportError("مكتبة openpyxl غير مثبتة. لقراءة ملفات Excel، يرجى تثبيتها باستخدام: pip install openpyxl")

    try:
        # Try with openpyxl first
        df = pd.read_excel(file, engine='openpyxl')
        return df
    except:
        try:
            # Fallback to default engine
            file.seek(0)
            df = pd.read_excel(file)
            return df
        except Exception as e:
            raise Exception(f"خطأ في قراءة ملف Excel: {str(e)}")

def validate_dataframe(df):
    """Validate the loaded dataframe"""
    issues = []

    if df is None or df.empty:
        issues.append("الملف لا يحتوي على بيانات")
        return False, issues

    if len(df.columns) == 0:
        issues.append("الملف لا يحتوي على أعمدة")
        return False, issues

    if len(df) < 3:
        issues.append("يحتاج التحليل إلى 3 مشاهدات على الأقل")
        return False, issues

    return True, issues

def ensure_numeric(df, columns):
    """Ensure columns are numeric for analysis"""
    numeric_df = df.copy()
    for col in columns:
        if col in numeric_df.columns:
            numeric_df[col] = pd.to_numeric(numeric_df[col], errors='coerce')
    return numeric_df

# ============================================
# Email sending function
# ============================================
def send_email_to_developer(name, message):
    """Send email to developer using mailto link"""
    try:
        # Save to session state
        st.session_state.contact_messages.append({
            'name': name,
            'message': message,
            'timestamp': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        })

        return True, "تم فتح بريدك الإلكتروني لإرسال الرسالة"

    except Exception as e:
        return False, f"حدث خطأ: {str(e)}"

# ============================================
# Show library availability warnings and installation instructions
# ============================================
with st.sidebar:
    st.markdown("### 📦 المكتبات المطلوبة")

    if not EXCEL_AVAILABLE:
        st.markdown("""
        <div class='install-box'>
            <strong>⚠️ دعم Excel غير متوفر</strong>
        </div>
        <div class='install-code'>
            pip install openpyxl
        </div>
        """, unsafe_allow_html=True)

    if not PLOTS_AVAILABLE:
        st.markdown("""
        <div class='install-box'>
            <strong>⚠️ الرسوم البيانية غير متوفرة</strong>
        </div>
        <div class='install-code'>
            pip install matplotlib seaborn
        </div>
        """, unsafe_allow_html=True)

    if not STATS_AVAILABLE:
        st.markdown("""
        <div class='install-box'>
            <strong>⚠️ الاختبارات الإحصائية غير متوفرة</strong>
        </div>
        <div class='install-code'>
            pip install scipy statsmodels
        </div>
        """, unsafe_allow_html=True)

    if not DOCX_AVAILABLE:
        st.markdown("""
        <div class='install-box'>
            <strong>⚠️ تصدير Word غير متوفر</strong>
        </div>
        <div class='install-code'>
            pip install python-docx
        </div>
        """, unsafe_allow_html=True)

# ============================================
# Sidebar - File Upload with Arabic text
# ============================================
with st.sidebar:
    st.markdown("---")
    st.markdown("""
    <div style='text-align: center; padding: 10px; background: linear-gradient(135deg, #1e3c72 0%, #2a5298 100%); border-radius: 10px; margin-bottom: 20px;'>
        <h3 style='color: white;'>📂 تحميل البيانات</h3>
    </div>
    """, unsafe_allow_html=True)

    # File uploader with Arabic text - modified the help text
    uploaded_file = st.file_uploader(
        "اختر ملف البيانات",
        type=['xlsx', 'xls', 'csv'],
        help="اسحب وأفلت الملف هنا"  # Changed to Arabic
    )

    # Show uploaded file info in Arabic
    if uploaded_file is not None:
        file_size = uploaded_file.size / 1024
        st.markdown(f"""
        <div class='uploadedFile'>
            <strong>الملف المرفوع:</strong> {uploaded_file.name}<br>
            <strong>الحجم:</strong> {file_size:.1f} كيلوبايت<br>
            <strong>النوع:</strong> {uploaded_file.type}
        </div>
        """, unsafe_allow_html=True)

    if st.button("🗑️ مسح الملف", use_container_width=True):
        st.session_state.data_loaded = False
        st.session_state.df = None
        st.session_state.show_results = False
        st.session_state.uploaded_filename = None
        st.rerun()

    if uploaded_file is not None:
        try:
            if st.session_state.uploaded_filename != uploaded_file.name:
                with st.spinner('جاري تحميل البيانات...'):
                    file_extension = uploaded_file.name.split('.')[-1].lower()

                    if file_extension == 'csv':
                        df, encoding = load_csv_with_encoding(uploaded_file)
                        st.success(f"✅ تم تحميل CSV بترميز {encoding}")
                    else:
                        # Check if Excel support is available
                        if not EXCEL_AVAILABLE:
                            st.error("❌ لا يمكن قراءة ملفات Excel. الرجاء تثبيت مكتبة openpyxl")
                            st.info("للتثبيت: pip install openpyxl")
                            st.stop()

                        df = load_excel_file(uploaded_file)
                        st.success("✅ تم تحميل Excel بنجاح")

                    is_valid, issues = validate_dataframe(df)

                    if is_valid:
                        st.session_state.df = df
                        st.session_state.data_loaded = True
                        st.session_state.uploaded_filename = uploaded_file.name

                        st.success(f"✅ {len(df)} سجل و {len(df.columns)} متغير")

                        with st.expander("🔍 معاينة البيانات"):
                            st.dataframe(df.head(5), use_container_width=True)
                    else:
                        for issue in issues:
                            st.error(issue)

        except ImportError as e:
            st.error(f"❌ {str(e)}")
            st.info("قم بتثبيت المكتبة المطلوبة باستخدام الأمر أعلاه")
        except Exception as e:
            st.error(f"❌ خطأ: {str(e)}")

# ============================================
# Likert scale analysis functions
# ============================================

def interpret_likert_trend(mean_score):
    """Interpret trend based on Likert scale (1-5)"""
    if pd.isna(mean_score):
        return {
            'trend': 'غير متوفر',
            'class': 'trend-moderate',
            'icon': '❓',
            'description': 'بيانات غير كافية',
            'color': '#9e9e9e'
        }
    if mean_score >= 4.5:
        return {
            'trend': 'مرتفع جداً',
            'class': 'trend-very-high',
            'icon': '🚀',
            'description': 'موافقة قوية جداً',
            'color': '#1b5e20'
        }
    elif mean_score >= 3.5:
        return {
            'trend': 'مرتفع',
            'class': 'trend-high',
            'icon': '📈',
            'description': 'موافقة',
            'color': '#2e7d32'
        }
    elif mean_score >= 2.5:
        return {
            'trend': 'متوسط',
            'class': 'trend-moderate',
            'icon': '📊',
            'description': 'محايد',
            'color': '#f9a825'
        }
    elif mean_score >= 1.5:
        return {
            'trend': 'منخفض',
            'class': 'trend-low',
            'icon': '📉',
            'description': 'عدم موافقة',
            'color': '#c62828'
        }
    else:
        return {
            'trend': 'منخفض جداً',
            'class': 'trend-very-low',
            'icon': '⚠️',
            'description': 'رفض قوي',
            'color': '#b71c1c'
        }

def calculate_likert_distribution(data):
    """Calculate distribution across Likert categories"""
    categories = {
        1: 'غير موافق بشدة',
        2: 'غير موافق',
        3: 'محايد',
        4: 'موافق',
        5: 'موافق بشدة'
    }

    distribution = {}
    data_clean = data.dropna()

    if len(data_clean) == 0:
        return distribution

    for value, label in categories.items():
        count = (data_clean.round() == value).sum()
        percentage = (count / len(data_clean)) * 100
        distribution[label] = {
            'count': count,
            'percentage': round(percentage, 1)
        }

    return distribution

def calculate_factor_trend(factor_data):
    """Calculate trend for a factor based on Likert scale"""
    results = {}

    try:
        data_clean = factor_data.dropna()

        if len(data_clean) == 0:
            return {
                'mean': 0,
                'median': 0,
                'std': 0,
                'trend': 'غير متوفر',
                'trend_class': 'trend-moderate',
                'icon': '❓',
                'description': 'بيانات غير كافية',
                'agreement_percent': 0,
                'disagreement_percent': 0,
                'neutral_percent': 0,
                'distribution': {},
                'n_cases': 0
            }

        mean_score = data_clean.mean()
        median_score = data_clean.median()
        std_dev = data_clean.std()

        # Get trend interpretation
        trend_info = interpret_likert_trend(mean_score)

        # Calculate distribution
        distribution = calculate_likert_distribution(data_clean)

        # Calculate agreement percentage (4+5)
        agreement = ((data_clean >= 4).sum() / len(data_clean)) * 100

        # Calculate disagreement percentage (1+2)
        disagreement = ((data_clean <= 2).sum() / len(data_clean)) * 100

        results = {
            'mean': round(mean_score, 2),
            'median': round(median_score, 2),
            'std': round(std_dev, 2),
            'trend': trend_info['trend'],
            'trend_class': trend_info['class'],
            'icon': trend_info['icon'],
            'description': trend_info['description'],
            'agreement_percent': round(agreement, 1),
            'disagreement_percent': round(disagreement, 1),
            'neutral_percent': round(100 - agreement - disagreement, 1),
            'distribution': distribution,
            'n_cases': len(data_clean)
        }

    except Exception as e:
        results = {
            'mean': 0,
            'median': 0,
            'std': 0,
            'trend': 'خطأ',
            'trend_class': 'trend-moderate',
            'icon': '❌',
            'description': str(e),
            'agreement_percent': 0,
            'disagreement_percent': 0,
            'neutral_percent': 0,
            'distribution': {},
            'n_cases': 0,
            'error': str(e)
        }

    return results

# ============================================
# Statistical test functions (with availability checks)
# ============================================

def interpret_p_value(p_value):
    """Interpret p-value"""
    if pd.isna(p_value):
        return "غير متاح"
    if p_value < 0.001:
        return "دال إحصائياً بمستوى مرتفع جداً (p < 0.001)"
    elif p_value < 0.01:
        return "دال إحصائياً بمستوى مرتفع (p < 0.01)"
    elif p_value < 0.05:
        return "دال إحصائياً (p < 0.05)"
    else:
        return "غير دال إحصائياً (p > 0.05)"

def check_normality(data):
    """Check normality using Shapiro-Wilk test"""
    if not STATS_AVAILABLE:
        return None
    try:
        data_clean = data.dropna()
        if len(data_clean) >= 3:
            statistic, p_value = shapiro(data_clean)
            return {
                'statistic': round(statistic, 4),
                'p_value': round(p_value, 4),
                'normal': p_value > 0.05,
                'interpretation': interpret_p_value(p_value)
            }
    except:
        pass
    return None

def perform_ttest(group1, group2, name1="المجموعة 1", name2="المجموعة 2"):
    """Perform independent t-test"""
    if not STATS_AVAILABLE:
        return {'error': 'المكتبات الإحصائية غير متوفرة'}

    results = {}

    try:
        g1 = group1.dropna()
        g2 = group2.dropna()

        if len(g1) < 2 or len(g2) < 2:
            results['error'] = "عدد المشاهدات غير كافٍ للاختبار"
            return results

        # Descriptive statistics
        results['descriptives'] = {
            name1: {
                'n': len(g1),
                'mean': round(g1.mean(), 3),
                'std': round(g1.std(), 3),
                'sem': round(g1.sem(), 3)
            },
            name2: {
                'n': len(g2),
                'mean': round(g2.mean(), 3),
                'std': round(g2.std(), 3),
                'sem': round(g2.sem(), 3)
            }
        }

        # Levene's test for homogeneity
        levene_stat, levene_p = levene(g1, g2)
        equal_var = levene_p > 0.05

        # T-test
        t_stat, t_p = ttest_ind(g1, g2, equal_var=equal_var)

        # Cohen's d effect size
        pooled_std = np.sqrt(((len(g1)-1)*g1.std()**2 + (len(g2)-1)*g2.std()**2) / (len(g1)+len(g2)-2))
        cohen_d = abs(g1.mean() - g2.mean()) / pooled_std if pooled_std != 0 else 0

        results['test'] = {
            't_statistic': round(t_stat, 3),
            'df': len(g1) + len(g2) - 2,
            'p_value': round(t_p, 4),
            'effect_size': round(cohen_d, 3),
            'significant': t_p < 0.05,
            'interpretation': interpret_p_value(t_p)
        }

        # Generate conclusion
        if t_p < 0.05:
            if g1.mean() > g2.mean():
                results['conclusion'] = f"توجد فروق دالة إحصائياً، حيث أن {name1} أعلى من {name2}"
            else:
                results['conclusion'] = f"توجد فروق دالة إحصائياً، حيث أن {name2} أعلى من {name1}"
        else:
            results['conclusion'] = "لا توجد فروق دالة إحصائياً بين المجموعتين"

    except Exception as e:
        results['error'] = str(e)

    return results

def perform_anova(groups, group_names):
    """Perform one-way ANOVA"""
    if not STATS_AVAILABLE:
        return {'error': 'المكتبات الإحصائية غير متوفرة'}

    results = {}

    try:
        clean_groups = [g.dropna() for g in groups]
        clean_groups = [g for g in clean_groups if len(g) >= 2]

        if len(clean_groups) < 2:
            results['error'] = "عدد المجموعات الصالحة غير كافٍ"
            return results

        # Descriptive statistics
        descriptives = []
        for name, group in zip(group_names, clean_groups):
            descriptives.append({
                'المجموعة': name,
                'العدد': len(group),
                'المتوسط': round(group.mean(), 3),
                'الانحراف المعياري': round(group.std(), 3)
            })
        results['descriptives'] = descriptives

        # Levene's test
        levene_stat, levene_p = levene(*clean_groups)
        results['homogeneity'] = {
            'levene_statistic': round(levene_stat, 3),
            'levene_p_value': round(levene_p, 4),
            'homogeneous': levene_p > 0.05
        }

        # ANOVA
        f_stat, f_p = f_oneway(*clean_groups)

        # Effect size (eta-squared)
        df_between = len(clean_groups) - 1
        df_within = sum(len(g) for g in clean_groups) - len(clean_groups)
        eta_sq = (f_stat * df_between) / (f_stat * df_between + df_within) if f_stat != 0 else 0

        results['test'] = {
            'f_statistic': round(f_stat, 3),
            'df_between': df_between,
            'df_within': df_within,
            'p_value': round(f_p, 4),
            'effect_size': round(eta_sq, 3),
            'significant': f_p < 0.05,
            'interpretation': interpret_p_value(f_p)
        }

        # Post-hoc if significant
        if f_p < 0.05 and len(clean_groups) > 2:
            data_for_tukey = []
            for name, group in zip(group_names, clean_groups):
                for val in group:
                    data_for_tukey.append({'group': name, 'value': val})

            tukey_df = pd.DataFrame(data_for_tukey)
            tukey_results = pairwise_tukeyhsd(tukey_df['value'], tukey_df['group'])

            posthoc = []
            for i in range(len(tukey_results.reject)):
                posthoc.append({
                    'المقارنة': f"{tukey_results.groupsunique[tukey_results._multicomp.pairindices[0][i]]} vs {tukey_results.groupsunique[tukey_results._multicomp.pairindices[1][i]]}",
                    'فرق المتوسطات': round(tukey_results.meandiffs[i], 3),
                    'p-value': round(tukey_results.pvalues[i], 4),
                    'الدلالة': 'دالة' if tukey_results.reject[i] else 'غير دالة'
                })
            results['posthoc'] = posthoc

        # Generate conclusion
        if f_p < 0.05:
            results['conclusion'] = "توجد فروق دالة إحصائياً بين المجموعات"
        else:
            results['conclusion'] = "لا توجد فروق دالة إحصائياً بين المجموعات"

    except Exception as e:
        results['error'] = str(e)

    return results

def perform_mannwhitney(group1, group2, name1="المجموعة 1", name2="المجموعة 2"):
    """Perform Mann-Whitney U test"""
    if not STATS_AVAILABLE:
        return {'error': 'المكتبات الإحصائية غير متوفرة'}

    results = {}

    try:
        g1 = group1.dropna()
        g2 = group2.dropna()

        if len(g1) < 2 or len(g2) < 2:
            results['error'] = "عدد المشاهدات غير كافٍ للاختبار"
            return results

        # Descriptive statistics
        results['descriptives'] = {
            name1: {
                'n': len(g1),
                'median': round(g1.median(), 3),
                'q1': round(g1.quantile(0.25), 3),
                'q3': round(g1.quantile(0.75), 3),
                'mean_rank': round(g1.rank().mean(), 2)
            },
            name2: {
                'n': len(g2),
                'median': round(g2.median(), 3),
                'q1': round(g2.quantile(0.25), 3),
                'q3': round(g2.quantile(0.75), 3),
                'mean_rank': round(g2.rank().mean(), 2)
            }
        }

        # Mann-Whitney U test
        u_stat, u_p = mannwhitneyu(g1, g2, alternative='two-sided')

        # Effect size (r)
        from scipy.stats import norm
        z_score = norm.ppf(u_p/2)
        n_total = len(g1) + len(g2)
        effect_size = abs(z_score) / np.sqrt(n_total) if z_score != 0 else 0

        results['test'] = {
            'u_statistic': round(u_stat, 3),
            'p_value': round(u_p, 4),
            'effect_size': round(effect_size, 3),
            'significant': u_p < 0.05,
            'interpretation': interpret_p_value(u_p)
        }

        # Generate conclusion
        if u_p < 0.05:
            if g1.rank().mean() > g2.rank().mean():
                results['conclusion'] = f"توجد فروق دالة إحصائياً، حيث أن {name1} أعلى من {name2}"
            else:
                results['conclusion'] = f"توجد فروق دالة إحصائياً، حيث أن {name2} أعلى من {name1}"
        else:
            results['conclusion'] = "لا توجد فروق دالة إحصائياً بين المجموعتين"

    except Exception as e:
        results['error'] = str(e)

    return results

def perform_kruskal_wallis(groups, group_names):
    """Perform Kruskal-Wallis test"""
    if not STATS_AVAILABLE:
        return {'error': 'المكتبات الإحصائية غير متوفرة'}

    results = {}

    try:
        clean_groups = [g.dropna() for g in groups]
        clean_groups = [g for g in clean_groups if len(g) >= 2]

        if len(clean_groups) < 2:
            results['error'] = "عدد المجموعات الصالحة غير كافٍ"
            return results

        # Descriptive statistics
        descriptives = []
        for name, group in zip(group_names, clean_groups):
            descriptives.append({
                'المجموعة': name,
                'العدد': len(group),
                'الوسيط': round(group.median(), 3),
                'متوسط الرتب': round(group.rank().mean(), 2)
            })
        results['descriptives'] = descriptives

        # Kruskal-Wallis test
        h_stat, h_p = kruskal(*clean_groups)

        results['test'] = {
            'h_statistic': round(h_stat, 3),
            'df': len(clean_groups) - 1,
            'p_value': round(h_p, 4),
            'significant': h_p < 0.05,
            'interpretation': interpret_p_value(h_p)
        }

        # Generate conclusion
        if h_p < 0.05:
            results['conclusion'] = "توجد فروق دالة إحصائياً بين المجموعات"
        else:
            results['conclusion'] = "لا توجد فروق دالة إحصائياً بين المجموعات"

    except Exception as e:
        results['error'] = str(e)

    return results

def perform_chi_square(observed, var1_name="", var2_name=""):
    """Perform Chi-square test"""
    if not STATS_AVAILABLE:
        return {'error': 'المكتبات الإحصائية غير متوفرة'}

    results = {}

    try:
        # Chi-square test
        chi2, p, dof, expected = chi2_contingency(observed)

        # Cramer's V effect size
        n = observed.sum().sum()
        min_dim = min(observed.shape) - 1
        cramer_v = np.sqrt(chi2 / (n * min_dim)) if min_dim > 0 and n > 0 else 0

        results['test'] = {
            'chi2': round(chi2, 3),
            'df': dof,
            'p_value': round(p, 4),
            'effect_size': round(cramer_v, 3),
            'significant': p < 0.05,
            'interpretation': interpret_p_value(p)
        }

        # Expected frequencies
        results['expected'] = expected

        # Generate conclusion
        if p < 0.05:
            results['conclusion'] = f"يوجد ارتباط دال إحصائياً بين {var1_name} و {var2_name}"
        else:
            results['conclusion'] = f"لا يوجد ارتباط دال إحصائياً بين {var1_name} و {var2_name}"

    except Exception as e:
        results['error'] = str(e)

    return results

def calculate_cronbach_alpha(df_items):
    """Calculate Cronbach's alpha"""
    try:
        # Convert to numeric and drop rows with any missing values
        df_numeric = df_items.apply(pd.to_numeric, errors='coerce')
        df_clean = df_numeric.dropna()

        if len(df_clean.columns) < 2 or len(df_clean) < 2:
            return None

        items = df_clean.values
        k = items.shape[1]
        item_variances = np.var(items, axis=0, ddof=1)
        total_scores = np.sum(items, axis=1)
        total_variance = np.var(total_scores, ddof=1)

        if total_variance > 0:
            alpha = (k / (k - 1)) * (1 - (np.sum(item_variances) / total_variance))

            # Interpret alpha
            if alpha >= 0.9:
                quality = "ممتاز"
            elif alpha >= 0.8:
                quality = "جيد"
            elif alpha >= 0.7:
                quality = "مقبول"
            elif alpha >= 0.6:
                quality = "ضعيف"
            else:
                quality = "غير مقبول"

            return {
                'alpha': round(alpha, 3),
                'quality': quality,
                'n_items': k,
                'n_cases': len(df_clean)
            }

        return None

    except Exception as e:
        return None

# ============================================
# Plotting functions (with French labels)
# ============================================

def create_likert_bar_chart(distribution, title):
    """Create bar chart for Likert distribution with French labels"""
    if not PLOTS_AVAILABLE or not distribution:
        return None

    try:
        fig, ax = plt.subplots(figsize=(10, 6))

        # French labels for Likert categories
        french_categories = {
            'غير موافق بشدة': 'Pas du tout d\'accord',
            'غير موافق': 'Pas d\'accord',
            'محايد': 'Neutre',
            'موافق': 'D\'accord',
            'موافق بشدة': 'Tout à fait d\'accord'
        }

        # Translate categories to French
        categories = list(distribution.keys())
        french_labels = [french_categories.get(cat, cat) for cat in categories]
        percentages = [distribution[cat]['percentage'] for cat in categories]
        colors = ['#d32f2f', '#f57c00', '#fdd835', '#7cb342', '#2e7d32']

        bars = ax.bar(range(len(categories)), percentages, color=colors, alpha=0.8)
        ax.set_title(title, fontsize=14, fontweight='bold', pad=20)
        ax.set_ylabel('Pourcentage (%)', fontsize=12)  # French label
        ax.set_ylim(0, 100)

        # Set French x-tick labels
        ax.set_xticks(range(len(categories)))
        ax.set_xticklabels(french_labels, rotation=45, ha='right')

        # Add percentage labels on bars
        for bar, pct in zip(bars, percentages):
            height = bar.get_height()
            ax.text(bar.get_x() + bar.get_width()/2., height,
                   f'{pct}%', ha='center', va='bottom', fontsize=10)

        plt.tight_layout()

        return fig
    except Exception as e:
        return None

def create_comparison_boxplot(groups, group_names, ylabel, title):
    """Create boxplot for group comparison with French labels"""
    if not PLOTS_AVAILABLE:
        return None

    try:
        fig, ax = plt.subplots(figsize=(10, 6))

        bp = ax.boxplot(groups, labels=group_names, patch_artist=True)
        colors = plt.cm.Set3(np.linspace(0, 1, len(groups)))

        for patch, color in zip(bp['boxes'], colors):
            patch.set_facecolor(color)

        ax.set_title(title, fontsize=14, fontweight='bold')
        ax.set_ylabel(ylabel, fontsize=12)
        ax.grid(True, alpha=0.3, axis='y')
        plt.xticks(rotation=45, ha='right')
        plt.tight_layout()

        return fig
    except Exception as e:
        return None

def create_correlation_heatmap(corr_matrix, title="Matrice de corrélation"):
    """Create correlation heatmap with French title"""
    if not PLOTS_AVAILABLE:
        return None

    try:
        fig, ax = plt.subplots(figsize=(10, 8))

        sns.heatmap(corr_matrix, annot=True, fmt='.3f', cmap='RdBu_r',
                   center=0, square=True, linewidths=0.5,
                   cbar_kws={"shrink": 0.8}, ax=ax)

        ax.set_title(title, fontsize=14, fontweight='bold', pad=20)
        ax.set_xlabel('Variables', fontsize=12)
        ax.set_ylabel('Variables', fontsize=12)
        plt.xticks(rotation=45, ha='right')
        plt.yticks(rotation=0)
        plt.tight_layout()

        return fig
    except Exception as e:
        return None

def create_normality_plot(data, var_name=""):
    """Create normality plots with French labels"""
    if not PLOTS_AVAILABLE:
        return None

    try:
        fig, axes = plt.subplots(2, 2, figsize=(12, 10))
        data_clean = data.dropna()

        # Histogram
        axes[0, 0].hist(data_clean, bins='auto', density=True, alpha=0.7,
                       color='steelblue', edgecolor='black')
        axes[0, 0].set_title(f'Distribution - {var_name}', fontsize=12, fontweight='bold')
        axes[0, 0].set_xlabel('Valeurs', fontsize=10)
        axes[0, 0].set_ylabel('Densité', fontsize=10)
        axes[0, 0].grid(True, alpha=0.3)

        # Q-Q plot
        stats.probplot(data_clean, dist="norm", plot=axes[0, 1])
        axes[0, 1].set_title('Graphique Q-Q', fontsize=12, fontweight='bold')
        axes[0, 1].set_xlabel('Quantiles théoriques', fontsize=10)
        axes[0, 1].set_ylabel('Quantiles observés', fontsize=10)
        axes[0, 1].grid(True, alpha=0.3)

        # Box plot
        axes[1, 0].boxplot(data_clean, vert=False, patch_artist=True)
        axes[1, 0].set_title('Diagramme en boîte', fontsize=12, fontweight='bold')
        axes[1, 0].set_xlabel('Valeurs', fontsize=10)
        axes[1, 0].grid(True, alpha=0.3)

        # Violin plot
        axes[1, 1].violinplot(data_clean, vert=False, showmeans=True, showmedians=True)
        axes[1, 1].set_title('Diagramme en violon', fontsize=12, fontweight='bold')
        axes[1, 1].set_xlabel('Valeurs', fontsize=10)
        axes[1, 1].grid(True, alpha=0.3)

        plt.suptitle(f'Analyse de la normalité - {var_name}', fontsize=14, fontweight='bold')
        plt.tight_layout()

        return fig
    except Exception as e:
        return None

def create_trend_plot(factor_data, factor_name, scale_min=1, scale_max=5):
    """Create trend visualization plot with French labels"""
    if not PLOTS_AVAILABLE:
        return None

    try:
        fig, axes = plt.subplots(1, 2, figsize=(14, 5))

        data_clean = factor_data.dropna()

        # Histogram with trend line
        axes[0].hist(data_clean, bins=15, alpha=0.7, color='steelblue',
                    edgecolor='black', density=True)
        axes[0].axvline(data_clean.mean(), color='red', linestyle='--',
                       linewidth=2, label=f'Moyenne: {data_clean.mean():.2f}')
        axes[0].axvline(scale_min + (scale_max - scale_min)/2, color='green',
                       linestyle=':', linewidth=2, label='Point milieu')
        axes[0].set_xlabel('Valeurs', fontsize=12)
        axes[0].set_ylabel('Densité', fontsize=12)
        axes[0].set_title(f'Distribution des scores - {factor_name}', fontsize=12, fontweight='bold')
        axes[0].legend()
        axes[0].grid(True, alpha=0.3)

        # Box plot with individual points
        bp = axes[1].boxplot(data_clean, vert=False, patch_artist=True, showmeans=True)
        bp['boxes'][0].set_facecolor('lightblue')

        # Add individual points with jitter
        y_jitter = np.random.normal(1, 0.05, size=len(data_clean))
        axes[1].scatter(data_clean, y_jitter, alpha=0.5, color='steelblue', s=30)

        axes[1].set_ylabel('')
        axes[1].set_xlabel('Valeurs', fontsize=12)
        axes[1].set_title(f'Diagramme de dispersion - {factor_name}', fontsize=12, fontweight='bold')
        axes[1].grid(True, alpha=0.3)

        plt.suptitle(f'Analyse de tendance - {factor_name}', fontsize=14, fontweight='bold')
        plt.tight_layout()

        return fig
    except Exception as e:
        return None

# ============================================
# Word report generation
# ============================================
def create_word_report(df, social_vars, factors, factor_trends):
    """Create Word report with analysis results"""
    if not DOCX_AVAILABLE:
        return None

    try:
        doc = Document()

        # Set document direction to RTL
        section = doc.sections[0]
        section.rtl = True

        # Title
        title = doc.add_heading('تقرير التحليل الإحصائي', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Date
        doc.add_paragraph(f'تاريخ التقرير: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        doc.add_paragraph(f'عدد الحالات: {len(df)}')
        doc.add_paragraph(f'عدد المتغيرات: {len(df.columns)}')

        # Social variables
        if social_vars:
            doc.add_heading('المتغيرات الاجتماعية', level=1)
            for var in social_vars:
                doc.add_heading(f'تحليل: {var}', level=2)
                freq = df[var].value_counts()
                for cat, count in freq.items():
                    percentage = (count / len(df)) * 100
                    doc.add_paragraph(f'{cat}: {count} ({percentage:.1f}%)')

        # Factors analysis
        if factors:
            doc.add_heading('تحليل المحاور', level=1)
            df_analysis = df.copy()

            # Convert to numeric for analysis
            for factor_name, items in factors.items():
                for item in items:
                    if item in df_analysis.columns:
                        df_analysis[item] = pd.to_numeric(df_analysis[item], errors='coerce')

            for factor_name, items in factors.items():
                # Select only numeric columns that exist
                valid_items = [item for item in items if item in df_analysis.columns]
                if valid_items:
                    df_analysis[factor_name] = df_analysis[valid_items].mean(axis=1)

            for factor_name, items in factors.items():
                doc.add_heading(f'المحور: {factor_name}', level=2)

                valid_items = [item for item in items if item in df_analysis.columns]
                if valid_items:
                    factor_data = df_analysis[valid_items].mean(axis=1).dropna()
                    trend = factor_trends.get(factor_name, {})

                    doc.add_paragraph(f'المتوسط: {factor_data.mean():.2f}' if len(factor_data) > 0 else 'المتوسط: غير متاح')
                    doc.add_paragraph(f'الوسيط: {factor_data.median():.2f}' if len(factor_data) > 0 else 'الوسيط: غير متاح')
                    doc.add_paragraph(f'الانحراف المعياري: {factor_data.std():.2f}' if len(factor_data) > 0 else 'الانحراف المعياري: غير متاح')
                    doc.add_paragraph(f'الاتجاه: {trend.get("trend", "غير محدد")}')
                    doc.add_paragraph(f'نسبة الموافقة: {trend.get("agreement_percent", 0)}%')

                    alpha = calculate_cronbach_alpha(df_analysis[valid_items])
                    if alpha:
                        doc.add_paragraph(f'معامل الثبات (ألفا كرونباخ): {alpha["alpha"]} - {alpha["quality"]}')

        # Save document
        doc_bytes = BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)
        return doc_bytes.getvalue()

    except Exception as e:
        st.error(f"خطأ في إنشاء تقرير Word: {str(e)}")
        return None

# ============================================
# Main content
# ============================================
if st.session_state.data_loaded:
    df = st.session_state.df.copy()
    all_columns = df.columns.tolist()

    # Create tabs
    tabs = st.tabs([
        "📋 تحديد المتغيرات",
        "📊 تحليل وصفي",
        "📈 تحليل الاتجاه",
        "🔬 اختبارات معلمية",
        "📉 اختبارات لا معلمية",
        "🔄 تحليل متقدم",
        "📥 تصدير النتائج",
        "📞 تواصل معنا"
    ])

    # ============================================
    # Tab 1: Variable Selection
    # ============================================
    with tabs[0]:
        st.markdown("""
        <div class='section-header'>
            <h3>🔍 تحديد المتغيرات الاجتماعية والمحاور</h3>
        </div>
        """, unsafe_allow_html=True)

        col1, col2 = st.columns([1, 3])

        with col1:
            st.markdown("### 🧑‍🤝‍🧑 المتغيرات الاجتماعية")
            social_vars = st.multiselect(
                "اختر المتغيرات الاجتماعية",
                options=all_columns
            )

        with col2:
            if social_vars:
                st.session_state.social_vars = social_vars
                st.success(f"✅ تم اختيار {len(social_vars)} متغير اجتماعي")

        if social_vars:
            st.markdown("---")
            st.markdown("### 📚 تحديد المحاور والفقرات")

            items_columns = [col for col in all_columns if col not in social_vars]

            if items_columns:
                n_factors = st.number_input(
                    "عدد المحاور",
                    min_value=1,
                    max_value=min(5, len(items_columns)),
                    value=1
                )

                factors = {}
                for i in range(n_factors):
                    with st.container():
                        st.markdown(f"<hr>", unsafe_allow_html=True)

                        col1, col2 = st.columns([1, 3])

                        with col1:
                            factor_name = st.text_input(
                                f"اسم المحور {i+1}",
                                value=f"المحور_{i+1}",
                                key=f"factor_name_{i}"
                            )

                        with col2:
                            factor_items = st.multiselect(
                                f"اختر فقرات المحور {i+1}",
                                options=items_columns,
                                key=f"factor_items_{i}"
                            )

                        if factor_items:
                            factors[factor_name] = factor_items
                            st.info(f"📌 {len(factor_items)} فقرة")

                if factors:
                    if st.button("🚀 بدء التحليل", type="primary", use_container_width=True):
                        st.session_state.factors = factors
                        st.session_state.show_results = True
                        st.success("✅ تم بدء التحليل بنجاح")
                        st.balloons()

    # ============================================
    # Tab 2: Descriptive Analysis
    # ============================================
    with tabs[1]:
        if st.session_state.show_results:
            st.markdown("""
            <div class='section-header'>
                <h3>📊 التحليل الوصفي</h3>
            </div>
            """, unsafe_allow_html=True)

            # Create a copy and ensure numeric columns for analysis
            df_analysis = df.copy()

            # Overall statistics
            st.markdown("### 📈 إحصائيات عامة")
            col1, col2, col3, col4 = st.columns(4)
            with col1:
                st.metric("إجمالي العينة", len(df))
            with col2:
                st.metric("عدد المتغيرات", len(df.columns))
            with col3:
                st.metric("عدد المحاور", len(st.session_state.factors))
            with col4:
                st.metric("المتغيرات الاجتماعية", len(st.session_state.social_vars))

            # Social variables statistics
            if st.session_state.social_vars:
                st.markdown("### 🧑‍🤝‍🧑 المتغيرات الاجتماعية")
                for var in st.session_state.social_vars:
                    if var in df.columns:
                        with st.expander(f"📌 تحليل: {var}"):
                            freq = df[var].value_counts().reset_index()
                            freq.columns = [var, 'التكرار']
                            freq['النسبة %'] = (freq['التكرار'] / len(df) * 100).round(2)
                            st.dataframe(freq, use_container_width=True)

            # Factors statistics
            if st.session_state.factors:
                st.markdown("### 📊 إحصائيات المحاور")
                factors_stats = []
                for factor_name, items in st.session_state.factors.items():
                    # Convert items to numeric and calculate mean
                    valid_items = []
                    for item in items:
                        if item in df.columns:
                            df_analysis[item] = pd.to_numeric(df_analysis[item], errors='coerce')
                            if not df_analysis[item].isna().all():
                                valid_items.append(item)

                    if valid_items:
                        factor_data = df_analysis[valid_items].mean(axis=1).dropna()
                        if len(factor_data) > 0:
                            factors_stats.append({
                                'المحور': factor_name,
                                'المتوسط': round(factor_data.mean(), 2),
                                'الوسيط': round(factor_data.median(), 2),
                                'الانحراف المعياري': round(factor_data.std(), 2),
                                'الحد الأدنى': round(factor_data.min(), 2),
                                'الحد الأقصى': round(factor_data.max(), 2),
                                'عدد الحالات': len(factor_data)
                            })

                if factors_stats:
                    st.dataframe(pd.DataFrame(factors_stats), use_container_width=True)
                else:
                    st.warning("⚠️ لا توجد بيانات رقمية كافية للتحليل")

    # ============================================
    # Tab 3: Trend Analysis (Likert-based)
    # ============================================
    with tabs[2]:
        if st.session_state.show_results:
            st.markdown("""
            <div class='section-header'>
                <h3>📈 تحليل الاتجاه وفق سلم ليكارت</h3>
            </div>
            """, unsafe_allow_html=True)

            # Create a copy and ensure numeric columns
            df_analysis = df.copy()

            # Likert scale explanation
            with st.expander("ℹ️ سلم ليكارت (1-5)"):
                st.markdown("""
                - **1**: غير موافق بشدة
                - **2**: غير موافق
                - **3**: محايد
                - **4**: موافق
                - **5**: موافق بشدة
                
                **تفسير المتوسطات:**
                - 4.5 - 5.0: موافقة قوية جداً (اتجاه مرتفع جداً)
                - 3.5 - 4.49: موافقة (اتجاه مرتفع)
                - 2.5 - 3.49: محايد (اتجاه متوسط)
                - 1.5 - 2.49: عدم موافقة (اتجاه منخفض)
                - 1.0 - 1.49: رفض قوي (اتجاه منخفض جداً)
                """)

            # Analyze each factor
            for factor_name, items in st.session_state.factors.items():
                st.markdown(f"### 📊 المحور: {factor_name}")

                # Convert items to numeric
                valid_items = []
                for item in items:
                    if item in df.columns:
                        df_analysis[item] = pd.to_numeric(df_analysis[item], errors='coerce')
                        if not df_analysis[item].isna().all():
                            valid_items.append(item)

                if valid_items:
                    factor_data = df_analysis[valid_items].mean(axis=1)

                    # Calculate trend
                    trend = calculate_factor_trend(factor_data)
                    st.session_state.factor_trends[factor_name] = trend

                    # Display trend
                    col1, col2 = st.columns([1, 1])

                    with col1:
                        st.markdown(f"""
                        <div class='{trend["trend_class"]}'>
                            <h3>{trend["icon"]} {trend["trend"]}</h3>
                            <p>{trend["description"]}</p>
                            <hr>
                            <p>المتوسط: {trend["mean"]}</p>
                            <p>الوسيط: {trend["median"]}</p>
                            <p>الانحراف: {trend["std"]}</p>
                            <p>ن = {trend["n_cases"]}</p>
                        </div>
                        """, unsafe_allow_html=True)

                    with col2:
                        st.markdown(f"""
                        <div style='background: white; padding: 15px; border-radius: 10px;'>
                            <h4>توزيع الاستجابات</h4>
                            <p style='color: #2e7d32;'>✅ موافق (4+5): {trend["agreement_percent"]}%</p>
                            <p style='color: #f9a825;'>➖ محايد (3): {trend["neutral_percent"]}%</p>
                            <p style='color: #c62828;'>❌ غير موافق (1+2): {trend["disagreement_percent"]}%</p>
                        </div>
                        """, unsafe_allow_html=True)

                    # Distribution chart (if available)
                    if PLOTS_AVAILABLE and trend['distribution']:
                        fig = create_likert_bar_chart(trend['distribution'], f'توزيع استجابات {factor_name}')
                        if fig:
                            st.pyplot(fig)
                            plt.close()

                    # Items analysis
                    st.markdown("#### 📝 تحليل الفقرات")
                    items_data = []
                    for item in valid_items:
                        item_data = df_analysis[item].dropna()
                        if len(item_data) > 0:
                            item_trend = interpret_likert_trend(item_data.mean())
                            items_data.append({
                                'الفقرة': item,
                                'المتوسط': round(item_data.mean(), 2),
                                'الانحراف': round(item_data.std(), 2),
                                'الاتجاه': item_trend['trend']
                            })

                    if items_data:
                        st.dataframe(pd.DataFrame(items_data), use_container_width=True)

                    # Reliability
                    alpha = calculate_cronbach_alpha(df_analysis[valid_items])
                    if alpha:
                        st.info(f"معامل الثبات (ألفا كرونباخ): {alpha['alpha']} - {alpha['quality']}")
                else:
                    st.warning(f"⚠️ لا توجد بيانات رقمية صالحة لتحليل المحور {factor_name}")

                st.markdown("---")

    # ============================================
    # Tab 4: Parametric Tests
    # ============================================
    with tabs[3]:
        if st.session_state.show_results:
            st.markdown("""
            <div class='section-header'>
                <h3>🔬 الاختبارات المعلمية</h3>
            </div>
            """, unsafe_allow_html=True)

            if not STATS_AVAILABLE:
                st.error("⚠️ المكتبات الإحصائية غير متوفرة. لا يمكن إجراء الاختبارات.")
                st.info("لتثبيت المكتبات: pip install scipy statsmodels")
            else:
                # Create a copy and ensure numeric columns
                df_analysis = df.copy()
                for factor_name, items in st.session_state.factors.items():
                    for item in items:
                        if item in df_analysis.columns:
                            df_analysis[item] = pd.to_numeric(df_analysis[item], errors='coerce')

                # Calculate factor scores
                for factor_name, items in st.session_state.factors.items():
                    valid_items = [item for item in items if item in df_analysis.columns]
                    if valid_items:
                        df_analysis[factor_name] = df_analysis[valid_items].mean(axis=1)

                if not st.session_state.factors:
                    st.warning("⚠️ الرجاء تحديد المحاور أولاً")
                elif not st.session_state.social_vars:
                    st.warning("⚠️ الرجاء تحديد المتغيرات الاجتماعية أولاً")
                else:
                    test_type = st.radio(
                        "اختر نوع التحليل",
                        ["اختبار T", "تحليل التباين (ANOVA)"],
                        horizontal=True
                    )

                    selected_factor = st.selectbox(
                        "اختر المحور",
                        options=list(st.session_state.factors.keys())
                    )

                    selected_var = st.selectbox(
                        "اختر المتغير الاجتماعي",
                        options=st.session_state.social_vars
                    )

                    if selected_var in df_analysis.columns:
                        groups = df_analysis[selected_var].dropna().unique()

                        if test_type == "اختبار T":
                            if len(groups) == 2:
                                group1 = df_analysis[df_analysis[selected_var] == groups[0]][selected_factor].dropna()
                                group2 = df_analysis[df_analysis[selected_var] == groups[1]][selected_factor].dropna()

                                if len(group1) >= 2 and len(group2) >= 2:
                                    results = perform_ttest(group1, group2, str(groups[0]), str(groups[1]))

                                    if 'error' not in results:
                                        col1, col2 = st.columns(2)
                                        with col1:
                                            st.info(f"**{groups[0]}**")
                                            st.write(f"ن = {results['descriptives'][str(groups[0])]['n']}")
                                            st.write(f"م = {results['descriptives'][str(groups[0])]['mean']}")
                                            st.write(f"ع = {results['descriptives'][str(groups[0])]['std']}")

                                        with col2:
                                            st.info(f"**{groups[1]}**")
                                            st.write(f"ن = {results['descriptives'][str(groups[1])]['n']}")
                                            st.write(f"م = {results['descriptives'][str(groups[1])]['mean']}")
                                            st.write(f"ع = {results['descriptives'][str(groups[1])]['std']}")

                                        st.markdown("#### 📊 نتائج الاختبار")
                                        test = results['test']

                                        col1, col2, col3, col4 = st.columns(4)
                                        with col1:
                                            st.metric("قيمة T", test['t_statistic'])
                                        with col2:
                                            st.metric("درجات الحرية", test['df'])
                                        with col3:
                                            st.metric("P-value", test['p_value'])
                                        with col4:
                                            if test['significant']:
                                                st.success("دال")
                                            else:
                                                st.warning("غير دال")

                                        st.info(f"📈 حجم التأثير (Cohen's d): {test['effect_size']}")
                                        st.success(f"💡 {results['conclusion']}")

                                        if PLOTS_AVAILABLE:
                                            fig = create_comparison_boxplot(
                                                [group1, group2],
                                                [str(groups[0]), str(groups[1])],
                                                selected_factor,
                                                f'مقارنة {selected_factor}'
                                            )
                                            if fig:
                                                st.pyplot(fig)
                                                plt.close()
                                    else:
                                        st.error(f"خطأ: {results['error']}")
                                else:
                                    st.warning("⚠️ عدد المشاهدات غير كافٍ في إحدى المجموعات")
                            else:
                                st.warning(f"⚠️ يحتاج اختبار T إلى متغير بفئتين")

                        else:  # ANOVA
                            if len(groups) >= 2:
                                groups_list = []
                                valid_groups = []
                                for group in groups:
                                    group_data = df_analysis[df_analysis[selected_var] == group][selected_factor].dropna()
                                    if len(group_data) >= 2:
                                        groups_list.append(group_data)
                                        valid_groups.append(str(group))

                                if len(groups_list) >= 2:
                                    results = perform_anova(groups_list, valid_groups)

                                    if 'error' not in results:
                                        st.dataframe(pd.DataFrame(results['descriptives']), use_container_width=True)

                                        test = results['test']

                                        col1, col2, col3, col4 = st.columns(4)
                                        with col1:
                                            st.metric("قيمة F", test['f_statistic'])
                                        with col2:
                                            st.metric("درجات الحرية", f"{test['df_between']},{test['df_within']}")
                                        with col3:
                                            st.metric("P-value", test['p_value'])
                                        with col4:
                                            if test['significant']:
                                                st.success("دال")
                                            else:
                                                st.warning("غير دال")

                                        st.info(f"📈 حجم التأثير (Eta-squared): {test['effect_size']}")
                                        st.success(f"💡 {results['conclusion']}")

                                        if 'posthoc' in results:
                                            st.markdown("#### 🔍 المقارنات البعدية")
                                            st.dataframe(pd.DataFrame(results['posthoc']), use_container_width=True)

                                        if PLOTS_AVAILABLE:
                                            fig = create_comparison_boxplot(
                                                groups_list,
                                                valid_groups,
                                                selected_factor,
                                                'مقارنة المجموعات'
                                            )
                                            if fig:
                                                st.pyplot(fig)
                                                plt.close()
                                    else:
                                        st.error(f"خطأ: {results['error']}")
                                else:
                                    st.warning("⚠️ لا توجد مجموعات ببيانات كافية للتحليل")
                            else:
                                st.warning(f"⚠️ المتغير {selected_var} لا يحتوي على مجموعات كافية")

    # ============================================
    # Tab 5: Non-parametric Tests
    # ============================================
    with tabs[4]:
        if st.session_state.show_results:
            st.markdown("""
            <div class='section-header'>
                <h3>📉 الاختبارات اللامعلمية</h3>
            </div>
            """, unsafe_allow_html=True)

            if not STATS_AVAILABLE:
                st.error("⚠️ المكتبات الإحصائية غير متوفرة. لا يمكن إجراء الاختبارات.")
                st.info("لتثبيت المكتبات: pip install scipy statsmodels")
            else:
                # Create a copy and ensure numeric columns
                df_analysis = df.copy()
                for factor_name, items in st.session_state.factors.items():
                    for item in items:
                        if item in df_analysis.columns:
                            df_analysis[item] = pd.to_numeric(df_analysis[item], errors='coerce')

                # Calculate factor scores
                for factor_name, items in st.session_state.factors.items():
                    valid_items = [item for item in items if item in df_analysis.columns]
                    if valid_items:
                        df_analysis[factor_name] = df_analysis[valid_items].mean(axis=1)

                if not st.session_state.factors:
                    st.warning("⚠️ الرجاء تحديد المحاور أولاً")
                elif not st.session_state.social_vars:
                    st.warning("⚠️ الرجاء تحديد المتغيرات الاجتماعية أولاً")
                else:
                    test_type = st.radio(
                        "اختر نوع الاختبار",
                        ["مان-ويتني", "كروسكال-واليس", "مربع كاي"],
                        horizontal=True
                    )

                    if test_type == "مان-ويتني":
                        selected_factor = st.selectbox(
                            "اختر المحور",
                            options=list(st.session_state.factors.keys()),
                            key="mw_factor"
                        )

                        selected_var = st.selectbox(
                            "اختر المتغير الاجتماعي",
                            options=st.session_state.social_vars,
                            key="mw_var"
                        )

                        if selected_var in df_analysis.columns:
                            groups = df_analysis[selected_var].unique()

                            if len(groups) == 2:
                                group1 = df_analysis[df_analysis[selected_var] == groups[0]][selected_factor].dropna()
                                group2 = df_analysis[df_analysis[selected_var] == groups[1]][selected_factor].dropna()

                                if len(group1) >= 2 and len(group2) >= 2:
                                    results = perform_mannwhitney(group1, group2, str(groups[0]), str(groups[1]))

                                    if 'error' not in results:
                                        col1, col2 = st.columns(2)
                                        with col1:
                                            st.info(f"**{groups[0]}**")
                                            st.write(f"ن = {results['descriptives'][str(groups[0])]['n']}")
                                            st.write(f"الوسيط = {results['descriptives'][str(groups[0])]['median']}")
                                            st.write(f"الربيعي = {results['descriptives'][str(groups[0])]['q1']} - {results['descriptives'][str(groups[0])]['q3']}")

                                        with col2:
                                            st.info(f"**{groups[1]}**")
                                            st.write(f"ن = {results['descriptives'][str(groups[1])]['n']}")
                                            st.write(f"الوسيط = {results['descriptives'][str(groups[1])]['median']}")
                                            st.write(f"الربيعي = {results['descriptives'][str(groups[1])]['q1']} - {results['descriptives'][str(groups[1])]['q3']}")

                                        test = results['test']
                                        col1, col2, col3 = st.columns(3)
                                        with col1:
                                            st.metric("قيمة U", test['u_statistic'])
                                        with col2:
                                            st.metric("P-value", test['p_value'])
                                        with col3:
                                            if test['significant']:
                                                st.success("دال")
                                            else:
                                                st.warning("غير دال")

                                        st.info(f"📈 حجم التأثير (r): {test['effect_size']}")
                                        st.success(f"💡 {results['conclusion']}")
                                    else:
                                        st.error(f"خطأ: {results['error']}")
                                else:
                                    st.warning("⚠️ عدد المشاهدات غير كافٍ في إحدى المجموعات")
                            else:
                                st.warning(f"⚠️ اختبار مان-ويتني يحتاج إلى متغير بفئتين")

                    elif test_type == "كروسكال-واليس":
                        selected_factor = st.selectbox(
                            "اختر المحور",
                            options=list(st.session_state.factors.keys()),
                            key="kw_factor"
                        )

                        selected_var = st.selectbox(
                            "اختر المتغير الاجتماعي",
                            options=st.session_state.social_vars,
                            key="kw_var"
                        )

                        if selected_var in df_analysis.columns:
                            groups = df_analysis[selected_var].unique()

                            if len(groups) >= 2:
                                groups_list = []
                                valid_groups = []
                                for group in groups:
                                    group_data = df_analysis[df_analysis[selected_var] == group][selected_factor].dropna()
                                    if len(group_data) >= 2:
                                        groups_list.append(group_data)
                                        valid_groups.append(str(group))

                                if len(groups_list) >= 2:
                                    results = perform_kruskal_wallis(groups_list, valid_groups)

                                    if 'error' not in results:
                                        st.dataframe(pd.DataFrame(results['descriptives']), use_container_width=True)

                                        test = results['test']
                                        col1, col2, col3 = st.columns(3)
                                        with col1:
                                            st.metric("قيمة H", test['h_statistic'])
                                        with col2:
                                            st.metric("درجات الحرية", test['df'])
                                        with col3:
                                            st.metric("P-value", test['p_value'])

                                        if test['significant']:
                                            st.success("✅ توجد فروق دالة")
                                        else:
                                            st.warning("⚠️ لا توجد فروق دالة")

                                        st.success(f"💡 {results['conclusion']}")
                                    else:
                                        st.error(f"خطأ: {results['error']}")
                                else:
                                    st.warning("⚠️ لا توجد مجموعات ببيانات كافية للتحليل")

                    else:  # Chi-square
                        if len(st.session_state.social_vars) >= 2:
                            var1 = st.selectbox(
                                "اختر المتغير الأول",
                                options=st.session_state.social_vars,
                                key="chi1"
                            )

                            var2 = st.selectbox(
                                "اختر المتغير الثاني",
                                options=[v for v in st.session_state.social_vars if v != var1],
                                key="chi2"
                            )

                            table = pd.crosstab(df[var1], df[var2])
                            st.write("**جدول التوافق:**")
                            st.dataframe(table, use_container_width=True)

                            results = perform_chi_square(table, var1, var2)

                            if 'error' not in results:
                                test = results['test']
                                col1, col2, col3, col4 = st.columns(4)
                                with col1:
                                    st.metric("مربع كاي", test['chi2'])
                                with col2:
                                    st.metric("درجات الحرية", test['df'])
                                with col3:
                                    st.metric("P-value", test['p_value'])
                                with col4:
                                    if test['significant']:
                                        st.success("دال")
                                    else:
                                        st.warning("غير دال")

                                st.info(f"📈 Cramer's V: {test['effect_size']}")
                                st.success(f"💡 {results['conclusion']}")

    # ============================================
    # Tab 6: Advanced Analysis
    # ============================================
    with tabs[5]:
        if st.session_state.show_results:
            st.markdown("""
            <div class='section-header'>
                <h3>🔄 تحليل متقدم</h3>
            </div>
            """, unsafe_allow_html=True)

            analysis_type = st.radio(
                "اختر نوع التحليل",
                ["📊 مصفوفة الارتباطات", "📈 الانحدار الخطي", "🔧 تحليل العوامل"],
                horizontal=True
            )

            # Create a copy and ensure numeric columns
            df_analysis = df.copy()
            for factor_name, items in st.session_state.factors.items():
                for item in items:
                    if item in df_analysis.columns:
                        df_analysis[item] = pd.to_numeric(df_analysis[item], errors='coerce')

            # Calculate factor scores
            for factor_name, items in st.session_state.factors.items():
                valid_items = [item for item in items if item in df_analysis.columns]
                if valid_items:
                    df_analysis[factor_name] = df_analysis[valid_items].mean(axis=1)

            if analysis_type == "📊 مصفوفة الارتباطات":
                if len(st.session_state.factors) >= 2:
                    # Create factors dataframe with only numeric columns
                    factors_data = {}
                    for name in st.session_state.factors.keys():
                        if name in df_analysis.columns:
                            factors_data[name] = df_analysis[name].dropna()

                    if len(factors_data) >= 2:
                        factors_df = pd.DataFrame(factors_data)

                        if STATS_AVAILABLE:
                            corr_matrix = factors_df.corr()
                            st.dataframe(corr_matrix.style.background_gradient(cmap='coolwarm').format("{:.3f}"),
                                       use_container_width=True)

                            if PLOTS_AVAILABLE:
                                fig, ax = plt.subplots(figsize=(10, 8))
                                sns.heatmap(corr_matrix, annot=True, fmt='.3f', cmap='RdBu_r', center=0, ax=ax)
                                ax.set_title('Matrice de corrélation')
                                st.pyplot(fig)
                                plt.close()
                        else:
                            st.warning("⚠️ المكتبات الإحصائية غير متوفرة لحساب الارتباطات")
                    else:
                        st.warning("⚠️ لا توجد بيانات كافية لحساب الارتباطات")

            elif analysis_type == "📈 الانحدار الخطي":
                if len(st.session_state.factors) >= 2 and STATS_AVAILABLE:
                    dependent = st.selectbox("المتغير التابع", options=list(st.session_state.factors.keys()))
                    independent = st.multiselect("المتغيرات المستقلة",
                                                [f for f in st.session_state.factors.keys() if f != dependent])

                    if independent and dependent in df_analysis.columns:
                        # Prepare data
                        valid_independent = [v for v in independent if v in df_analysis.columns]
                        if valid_independent:
                            data_for_reg = df_analysis[valid_independent + [dependent]].dropna()

                            if len(data_for_reg) > len(valid_independent) + 1:
                                X = data_for_reg[valid_independent]
                                y = data_for_reg[dependent]
                                X = sm.add_constant(X)

                                model = sm.OLS(y, X).fit()

                                st.write(f"**R-squared:** {model.rsquared:.3f}")
                                st.write(f"**R-squared المعدل:** {model.rsquared_adj:.3f}")
                                st.write(f"**F-statistic:** {model.fvalue:.3f}")
                                st.write(f"**P-value (F):** {model.f_pvalue:.4f}")

                                coef_df = pd.DataFrame({
                                    'المتغير': model.params.index,
                                    'المعامل': model.params.values.round(3),
                                    'P-value': model.pvalues.values.round(4)
                                })
                                st.dataframe(coef_df, use_container_width=True)
                            else:
                                st.warning("⚠️ بيانات غير كافية للانحدار الخطي")
                else:
                    if not STATS_AVAILABLE:
                        st.warning("⚠️ المكتبات الإحصائية غير متوفرة للانحدار الخطي")

            else:  # Factor Analysis
                st.info("🔧 تحليل العوامل يتطلب مكتبة scikit-learn")
                try:
                    from sklearn.decomposition import FactorAnalysis
                    from sklearn.preprocessing import StandardScaler

                    all_items = []
                    for items in st.session_state.factors.values():
                        all_items.extend(items)

                    if len(all_items) >= 3:
                        # Convert to numeric
                        df_factors = df[all_items].apply(pd.to_numeric, errors='coerce').dropna()

                        if len(df_factors) > 10:  # Need sufficient sample
                            scaler = StandardScaler()
                            df_scaled = scaler.fit_transform(df_factors)

                            n_factors = st.slider("عدد العوامل", 1, min(5, len(all_items)), min(3, len(all_items)))

                            fa = FactorAnalysis(n_components=n_factors, random_state=42)
                            factors = fa.fit_transform(df_scaled)

                            loadings = pd.DataFrame(
                                fa.components_.T,
                                index=all_items,
                                columns=[f'عامل{i+1}' for i in range(n_factors)]
                            )

                            st.dataframe(loadings.style.background_gradient(cmap='coolwarm').format("{:.3f}"),
                                       use_container_width=True)
                        else:
                            st.warning("⚠️ البيانات غير كافية لتحليل العوامل")
                    else:
                        st.warning("⚠️ تحتاج إلى 3 فقرات على الأقل لتحليل العوامل")

                except ImportError:
                    st.info("لتشغيل تحليل العوامل: قم بتثبيت scikit-learn: pip install scikit-learn")

    # ============================================
    # Tab 7: Export Results
    # ============================================
    with tabs[6]:
        if st.session_state.show_results:
            st.markdown("""
            <div class='section-header'>
                <h3>📥 تصدير النتائج</h3>
            </div>
            """, unsafe_allow_html=True)

            export_format = st.radio(
                "اختر صيغة التصدير",
                ["📄 CSV", "📊 Excel", "📝 Word (تقرير كامل)"],
                horizontal=True
            )

            if st.button("🔄 تجهيز النتائج", type="primary"):
                with st.spinner('جاري التجهيز...'):

                    if export_format == "📄 CSV":
                        # Prepare results for CSV
                        results_data = []

                        # Social variables
                        for var in st.session_state.social_vars:
                            if var in df.columns:
                                freq = df[var].value_counts()
                                for cat, count in freq.items():
                                    results_data.append({
                                        'التحليل': 'إحصائيات وصفية',
                                        'المتغير': var,
                                        'الفئة': str(cat),
                                        'التكرار': count,
                                        'النسبة': round(count/len(df)*100, 2)
                                    })

                        # Factors
                        df_analysis = df.copy()
                        for factor_name, items in st.session_state.factors.items():
                            valid_items = []
                            for item in items:
                                if item in df_analysis.columns:
                                    df_analysis[item] = pd.to_numeric(df_analysis[item], errors='coerce')
                                    if not df_analysis[item].isna().all():
                                        valid_items.append(item)

                            if valid_items:
                                factor_score = df_analysis[valid_items].mean(axis=1)
                                trend = st.session_state.factor_trends.get(factor_name, {})

                                results_data.append({
                                    'التحليل': 'تحليل المحاور',
                                    'المتغير': factor_name,
                                    'المتوسط': round(factor_score.mean(), 2) if len(factor_score) > 0 else 0,
                                    'الانحراف': round(factor_score.std(), 2) if len(factor_score) > 0 else 0,
                                    'الاتجاه': trend.get('trend', ''),
                                    'موافق%': trend.get('agreement_percent', 0)
                                })

                        if results_data:
                            results_df = pd.DataFrame(results_data)
                            csv = results_df.to_csv(index=False).encode('utf-8-sig')
                            st.download_button(
                                "📥 تحميل CSV",
                                csv,
                                f"نتائج_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv",
                                "text/csv",
                                use_container_width=True
                            )
                        else:
                            st.warning("⚠️ لا توجد نتائج للتصدير")

                    elif export_format == "📊 Excel":
                        # Prepare Excel file
                        output = BytesIO()
                        with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
                            # Social variables sheet
                            social_data = []
                            for var in st.session_state.social_vars:
                                if var in df.columns:
                                    freq = df[var].value_counts()
                                    for cat, count in freq.items():
                                        social_data.append({
                                            'المتغير': var,
                                            'الفئة': str(cat),
                                            'التكرار': count,
                                            'النسبة': round(count/len(df)*100, 2)
                                        })
                            if social_data:
                                pd.DataFrame(social_data).to_excel(writer, sheet_name='المتغيرات_الاجتماعية', index=False)

                            # Factors sheet
                            factors_data = []
                            df_analysis = df.copy()
                            for factor_name, items in st.session_state.factors.items():
                                valid_items = []
                                for item in items:
                                    if item in df_analysis.columns:
                                        df_analysis[item] = pd.to_numeric(df_analysis[item], errors='coerce')
                                        if not df_analysis[item].isna().all():
                                            valid_items.append(item)

                                if valid_items:
                                    factor_score = df_analysis[valid_items].mean(axis=1)
                                    trend = st.session_state.factor_trends.get(factor_name, {})
                                    factors_data.append({
                                        'المحور': factor_name,
                                        'المتوسط': round(factor_score.mean(), 2) if len(factor_score) > 0 else 0,
                                        'الانحراف': round(factor_score.std(), 2) if len(factor_score) > 0 else 0,
                                        'الاتجاه': trend.get('trend', ''),
                                        'موافق%': trend.get('agreement_percent', 0)
                                    })

                            if factors_data:
                                pd.DataFrame(factors_data).to_excel(writer, sheet_name='تحليل_المحاور', index=False)

                            # Correlations sheet (if available)
                            if len(st.session_state.factors) >= 2 and STATS_AVAILABLE:
                                factors_corr = {}
                                for name in st.session_state.factors.keys():
                                    if name in df_analysis.columns:
                                        factors_corr[name] = df_analysis[name].dropna()

                                if len(factors_corr) >= 2:
                                    factors_df = pd.DataFrame(factors_corr)
                                    factors_df.corr().to_excel(writer, sheet_name='الارتباطات')

                            # Info sheet
                            pd.DataFrame({
                                'معلومة': ['عدد الحالات', 'عدد المتغيرات', 'تاريخ التقرير'],
                                'القيمة': [len(df), len(df.columns), datetime.now().strftime('%Y-%m-%d')]
                            }).to_excel(writer, sheet_name='معلومات', index=False)

                        st.download_button(
                            "📥 تحميل Excel",
                            output.getvalue(),
                            f"نتائج_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx",
                            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                            use_container_width=True
                        )

                    else:  # Word
                        if DOCX_AVAILABLE:
                            word_data = create_word_report(
                                df,
                                st.session_state.social_vars,
                                st.session_state.factors,
                                st.session_state.factor_trends
                            )
                            if word_data:
                                st.download_button(
                                    "📥 تحميل تقرير Word",
                                    word_data,
                                    f"تقرير_تحليل_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                    use_container_width=True
                                )
                            else:
                                st.error("❌ خطأ في إنشاء تقرير Word")
                        else:
                            st.error("❌ مكتبة python-docx غير مثبتة. لتثبيتها: pip install python-docx")

                    st.success("✅ تم تجهيز الملف")

    # ============================================
    # Tab 8: Contact - Simplified with name and message
    # ============================================
    with tabs[7]:
        st.markdown("""
        <div class='contact-info'>
            <h3>📬 تواصل معنا</h3>
            <p>مرحباً بمقترحاتكم أو نصائحكم لتطوير التطبيق</p>
        </div>
        """, unsafe_allow_html=True)

        with st.form("contact_form"):
            name = st.text_input("الاسم واللقب", placeholder="أدخل اسمك الكامل")
            message = st.text_area("نص الرسالة", placeholder="اكتب مقترحاتك أو نصائحك هنا...", height=150)

            col1, col2, col3 = st.columns(3)
            with col2:
                submitted = st.form_submit_button("📨 إرسال", type="primary", use_container_width=True)

            if submitted:
                if not name or not message:
                    st.warning("⚠️ الرجاء ملء جميع الحقول")
                else:
                    # Create mailto link with the message
                    subject = f"مقترح من {name}"
                    body = f"الاسم: {name}\n\nالرسالة:\n{message}"
                    mailto_link = f"mailto:boutoubaamed@gmail.com?subject={subject}&body={body}"

                    # Create an HTML link that will open the user's email client
                    email_button = f"""
                    <div style='text-align: center; margin: 20px 0;'>
                        <a href='{mailto_link}' target='_blank' style='background: linear-gradient(135deg, #ffd700 0%, #ffb300 100%); color: #1e3c72; padding: 15px 30px; border-radius: 50px; text-decoration: none; font-weight: bold; font-size: 1.2em; display: inline-block;'>
                            📧 اضغط هنا لفتح بريدك الإلكتروني
                        </a>
                    </div>
                    """
                    st.markdown(email_button, unsafe_allow_html=True)
                    st.success("✅ تم تجهيز الرسالة. اضغط على الزر أعلاه لإرسالها عبر بريدك الإلكتروني.")
                    st.balloons()

# ============================================
# Welcome message for new users
# ============================================
if not st.session_state.data_loaded:
    st.markdown("""
    <div style='text-align: center; padding: 50px; background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); border-radius: 15px; color: white; margin: 20px 0;'>
        <h2 style='color: white; font-size: 2.5em; margin-bottom: 20px;'>🌟 مرحباً بك في منصة التحليل الإحصائي المتقدم</h2>
        <p style='font-size: 1.3em; margin-bottom: 30px;'>قم بتحميل بياناتك للبدء في التحليل الإحصائي المتكامل</p>
        <div style='display: flex; justify-content: center; gap: 20px; flex-wrap: wrap;'>
            <div style='background: rgba(255,255,255,0.2); padding: 15px; border-radius: 10px;'>
                <h3>📊 تحليل وصفي</h3>
            </div>
            <div style='background: rgba(255,255,255,0.2); padding: 15px; border-radius: 10px;'>
                <h3>📈 تحليل الاتجاه</h3>
            </div>
            <div style='background: rgba(255,255,255,0.2); padding: 15px; border-radius: 10px;'>
                <h3>🔬 اختبارات معلمية</h3>
            </div>
            <div style='background: rgba(255,255,255,0.2); padding: 15px; border-radius: 10px;'>
                <h3>📉 اختبارات لا معلمية</h3>
            </div>
        </div>
    </div>
    """, unsafe_allow_html=True)

# ============================================
# Footer with updated version
# ============================================
st.markdown("---")

footer_col1, footer_col2, footer_col3 = st.columns([1, 2, 1])

with footer_col1:
    st.markdown("""
    <div style='text-align: center; padding: 10px;'>
        <p>📊 إصدار تجريبي  ©  2026</p>
    </div>
    """, unsafe_allow_html=True)

with footer_col2:
    st.markdown("""
    <div style='text-align: center; padding: 10px;'>
        <p> جامعة عين تموشنت </p>
        <p style='font-size: 0.8em; opacity: 0.7;'>Developed by Pr Mohammed Boutouba</p>
    </div>
    """, unsafe_allow_html=True)

with footer_col3:
    st.markdown("""
    <div style='text-align: center; padding: 10px;'>
        <a href='mailto:boutoubaamed@gmail.com' style='margin: 0 5px; text-decoration: none; font-size: 24px;' title='إرسال بريد للمطور'>📧</a>
    </div>
    """, unsafe_allow_html=True)

# ============================================
# Main function
# ============================================
def main():
    """Main application function"""
    try:
        if st.session_state.data_loaded and st.session_state.show_results:
            if not st.session_state.social_vars:
                st.sidebar.info("ℹ️ يرجى تحديد المتغيرات الاجتماعية في تبويب 'تحديد المتغيرات'")
            if not st.session_state.factors:
                st.sidebar.info("ℹ️ يرجى تحديد المحاور في تبويب 'تحديد المتغيرات'")
    except Exception as e:
        st.error(f"❌ حدث خطأ: {str(e)}")

if __name__ == "__main__":
    main()